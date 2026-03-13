import streamlit as st
import fitz
import re
from docx import Document
import tempfile

st.title("Credit Repair Packet Generator")

uploaded_file = st.file_uploader("Upload Credit Report PDF", type="pdf")

accounts = []

def extract_accounts(pdf):
    accounts = []
    doc = fitz.open(stream=pdf.read(), filetype="pdf")

    text = ""
    for page in doc:
        text += page.get_text()

    # find creditor names
    creditors = re.findall(r'[A-Z][A-Z\s]+(?:SERV|BANK|FINANCE|CREDIT|AUTO|CAPITAL|MIDLAND|AFFIRM|CONN)', text)

    balances = re.findall(r'Balance\s*\$?([0-9,]+)', text)

    for i in range(min(len(creditors), len(balances))):
        accounts.append({
            "creditor": creditors[i],
            "balance": int(balances[i].replace(",", "")),
        })

    return accounts


def create_packet(accounts):

    doc = Document()
    doc.add_heading("Credit Repair Packet", level=1)

    doc.add_heading("Action Plan", level=2)

    for acc in accounts:

        action = "623 Dispute" if acc["balance"] > 1000 else "Pay For Delete"

        doc.add_paragraph(
            f"{acc['creditor']} - ${acc['balance']} - {action}"
        )

    doc.add_heading("Letters", level=2)

    for acc in accounts:

        if acc["balance"] > 1000:

            doc.add_heading(acc["creditor"], level=3)

            doc.add_paragraph("""
Date:

Re: Direct Dispute Under FCRA Section 623

To Whom It May Concern,

I am disputing the accuracy of the account referenced above being reported on my credit file.

Please provide documentation verifying this debt including:

1. Original signed contract
2. Complete payment history
3. Proof of ownership of the debt
4. Verification of date of first delinquency

If this information cannot be verified, the Fair Credit Reporting Act requires the account be deleted.

Sincerely,
Borrower
""")

        else:

            doc.add_heading(acc["creditor"], level=3)

            doc.add_paragraph("""
Date:

Re: Pay For Delete Settlement Offer

To Whom It May Concern,

I am willing to resolve the above account with a settlement payment if your company agrees to delete the tradeline from all credit bureaus.

Please confirm deletion agreement in writing.

Sincerely,
Borrower
""")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)

    return temp.name


if uploaded_file:

    accounts = extract_accounts(uploaded_file)

    st.write("Accounts Found")

    st.write(accounts)

    if st.button("Generate Credit Repair Packet"):

        file = create_packet(accounts)

        with open(file, "rb") as f:
            st.download_button(
                "Download Packet",
                f,
                file_name="credit_repair_packet.docx"
            )
